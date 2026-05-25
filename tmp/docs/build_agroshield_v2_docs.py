from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\jose\Documents\CopernicusLAC")
JIC_OUT = ROOT / "AgroShield_Propuesta_Tecnica_JIC_UTP_v2.docx"
ROADMAP_OUT = ROOT / "AgroShield_Roadmap_Arquitectura_Detallada.docx"


GREEN = "D9EAD3"
BLUE = "D9EAF7"
YELLOW = "FFF2CC"
RED = "F4CCCC"
GRAY = "EDEDED"
INK = (32, 32, 32)
ACCENT = (31, 78, 121)


def set_run_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_page(section, two_cols=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    if two_cols:
        set_columns(section, 2)


def set_columns(section, num=2, space_twips=567):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if cols.getparent() is None:
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=65, start=90, bottom=65, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=8.0, bold=False, italic=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(cell)
    if fill:
        cell_shading(cell, fill)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def compact_table(doc, rows, widths=None, header_fill=BLUE, font_size=7.7):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    if widths:
        set_table_width(table, widths)
    for idx, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, size=font_size, bold=True, fill=header_fill, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
    return table


def para(doc, text="", lead=None, size=10, first_indent=True, justify=True, after=0, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.0
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r = p.add_run(lead + " ")
        set_run_font(r, size=size, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def heading(doc, number, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number} {title}")
    set_run_font(r, size=14 if level == 1 else 10, bold=True, color=ACCENT if level == 1 else INK)
    return p


def unnumbered(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    set_run_font(r, size=14, bold=True, color=ACCENT)
    return p


def caption(doc, text, kind="table"):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "table" else WD_ALIGN_PARAGRAPH.JUSTIFY
    first, rest = text.split(" ", 1)
    r1 = p.add_run(first + " ")
    set_run_font(r1, size=8, bold=True)
    r2 = p.add_run(rest)
    set_run_font(r2, size=8)


def source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("Fuente: ")
    set_run_font(r, size=8, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_base_styles(doc):
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=run.font.size.pt if run.font.size else 10)


def build_jic_v2():
    doc = Document()
    set_page(doc.sections[0])
    set_base_styles(doc)

    p = doc.add_paragraph()
    r = p.add_run("Categoría prototipado: seguridad alimentaria y agricultura sostenible con observación de la Tierra.")
    set_run_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("AgroShield: semáforo Copernicus de riesgo hídrico para inocuidad agrícola en Azuero")
    set_run_font(r, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("AgroShield: Copernicus water-risk traffic light for agricultural food safety in Azuero")
    set_run_font(r, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("José Rodríguez1*, Yanesis Valdés2")
    set_run_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("1Afiliación pendiente; 2Afiliación pendiente")
    set_run_font(r, size=10, italic=True)

    para(
        doc,
        "En Azuero, el agua del río La Villa sostiene riego, lavado, producción agroexportadora y continuidad alimentaria. Cuando eventos de lluvia movilizan sedimento y contaminantes desde fuentes difusas, productores y técnicos carecen de una señal temprana, reproducible y trazable para decidir bajo incertidumbre. AgroShield propone un semáforo de inocuidad hídrica agrícola basado en Copernicus: Sentinel-2 vía CDSE para comparar señales ópticas sobre el área Chitré-La Arena, Sentinel-1 y CHIRPS como soporte ante lluvia y nubosidad, y un reporte técnico simple que explique estado, evidencia, confianza y límites. El prototipo no detecta pesticidas disueltos ni sustituye laboratorio; transforma observación satelital en evidencia preventiva para una cuenca real, una crisis real y una decisión agrícola concreta.",
        lead="Resumen",
        first_indent=False,
    )
    para(doc, "AgroShield, Copernicus, inocuidad agrícola, Río La Villa, Sentinel.", lead="Palabras clave", first_indent=False)
    para(
        doc,
        "In Azuero, the La Villa river sustains irrigation, washing, agro-export activity and food continuity. When rainfall events mobilize sediment and diffuse contaminants, producers and watershed technicians lack an early, reproducible and traceable signal for decision-making under uncertainty. AgroShield proposes a Copernicus-based agricultural water-safety traffic light: Sentinel-2 through CDSE to compare optical evidence over the Chitré-La Arena area, Sentinel-1 and CHIRPS as support under rainfall and cloud conditions, and a simple technical report explaining status, evidence, confidence and limitations.",
        lead="Abstract",
        first_indent=False,
    )
    para(doc, "AgroShield, Copernicus, food safety, La Villa river, Sentinel.", lead="Keywords", first_indent=False)
    p = doc.add_paragraph()
    r = p.add_run("* Corresponding author: jose.rodriguez104@utp.ac.pa")
    set_run_font(r, size=8)

    body = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(body, two_cols=True)

    heading(doc, "1.", "Introducción")
    para(
        doc,
        "La propuesta parte de una premisa operacional: si el río falla, falla la agricultura que depende de él. En la península de Azuero, la cuenca del río La Villa vincula abastecimiento, riego, inocuidad poscosecha y confianza en cadenas de exportación. Un técnico de cuenca necesita saber cuándo una señal hidrosedimentaria merece atención antes de que el problema se convierta en pérdida de cultivo, suspensión de riego o crisis pública [1], [9], [10]."
    )
    para(
        doc,
        "Las soluciones tradicionales aportan mediciones puntuales o inspecciones posteriores. Copernicus agrega algo distinto: visión sinóptica, archivo histórico y bajo costo marginal para mirar la cuenca completa [1]-[4]. AgroShield no promete una verdad química desde el espacio; propone una herramienta de priorización que responde una pregunta de hackathon: ¿puede una señal Sentinel reproducible ayudar a decidir riesgo hídrico agrícola en Azuero?"
    )

    heading(doc, "2.", "Diseño y metodología")
    para(
        doc,
        "El MVP se reduce a cuatro piezas: un área de interés real en Chitré-La Arena, una serie histórica Sentinel-2 consultada por CDSE, reglas transparentes de semáforo y un reporte técnico. Sentinel-1 y CHIRPS se conservan como soporte cuando la lluvia o nubosidad explican la incertidumbre; no se incluyen módulos agronómicos adicionales en el cuerpo principal."
    )
    caption(doc, "Figura 1. Flujo mínimo del MVP AgroShield", kind="figure")
    compact_table(
        doc,
        [
            ["AOI", "CDSE", "Señal", "Decisión"],
            ["Chitré-La Arena", "Sentinel-2 L2A", "Agua + sedimento", "Semáforo"],
            ["Misma geometría", "Serie histórica", "Baseline vs. evento", "Reporte técnico"],
        ],
        [Cm(1.8), Cm(1.9), Cm(2.2), Cm(1.8)],
        header_fill=GREEN,
        font_size=7.4,
    )
    source(doc, "Elaboración propia. El MVP se evalúa por reproducibilidad de la señal, no por cantidad de módulos.")

    heading(doc, "2.1", "Reglas auditables del semáforo", level=2)
    para(
        doc,
        "La primera versión evita pesos arbitrarios. El semáforo usa reglas explicables: si no hay suficientes píxeles válidos, el estado es incierto; si la máscara de agua confirma el río y una señal de turbidez aumenta frente a la línea base, el estado pasa a amarillo o rojo; si NDCI sube, se activa una bandera secundaria de posible carga biológica; si CHIRPS registra lluvia antecedente alta, aumenta la prioridad de revisión. Cada salida debe mostrar datos, fecha y confianza."
    )
    caption(doc, "Tabla 1. Reglas mínimas para clasificar evidencia")
    compact_table(
        doc,
        [
            ["Condición", "Salida"],
            ["Pocos píxeles válidos o nube dominante", "Incierto; no forzar verde."],
            ["Agua confirmada y señal óptica cercana a baseline", "Verde técnico."],
            ["Anomalía óptica moderada frente a baseline", "Amarillo; revisar riego."],
            ["Anomalía fuerte + lluvia antecedente", "Rojo; acción preventiva."],
            ["NDCI alto", "Bandera secundaria, no diagnóstico químico."],
        ],
        [Cm(3.0), Cm(4.7)],
        header_fill=YELLOW,
        font_size=7.5,
    )
    source(doc, "Reglas preliminares sujetas a calibración con datos históricos y observación local.")

    heading(doc, "3.", "Prueba central del MVP")
    para(
        doc,
        "El MVP será juzgado por una prueba reina: comparar la misma área de interés en tres contextos de 2025. Si CDSE devuelve una señal Sentinel reproducible sobre Chitré-La Arena y esa señal se puede explicar con incertidumbre explícita, AgroShield tiene base técnica para avanzar. Si la señal falla por geometría o nubosidad, el proyecto pivota a un tramo más amplio o a atlas de riesgo de cuenca."
    )
    caption(doc, "Tabla 2. Matriz de validación histórica")
    compact_table(
        doc,
        [
            ["Ventana", "Propósito", "Evidencia esperada"],
            ["Crisis 2025", "Buscar firma hidrosedimentaria.", "Cambio óptico o incertidumbre documentada."],
            ["Baseline cercana", "Comparar la misma estación.", "Valores normales sobre el mismo AOI."],
            ["Control estacional", "Evitar falsa narrativa.", "Separar temporada de evento."],
        ],
        [Cm(1.9), Cm(2.7), Cm(3.1)],
        header_fill=BLUE,
        font_size=7.5,
    )
    source(doc, "Diseño de verificación interna previo a la implementación formal del hackathon.")

    heading(doc, "4.", "Construcción y validación del prototipo")
    para(
        doc,
        "La construcción inicia con código local de verificación y datos cacheados. El resultado demostrable no es una suite completa, sino una pantalla o reporte con mapa del AOI, tabla de fechas, señal óptica, confianza y recomendación. El reporte puede exportarse como PDF o HTML; la interpretación asistida por IA se limita a una plantilla generada a partir de valores calculados, sin inferencias externas ni datos inventados."
    )
    para(
        doc,
        "El bot comunitario se mantiene como canal secundario: consulta de estado, recepción de reporte y posibilidad futura de registrar observaciones de campo. NitroSync y ERA5-Land quedan como trabajo futuro para ventanas de manejo agronómico, porque requieren cultivo, etapa fenológica y calibración local."
    )

    heading(doc, "5.", "Oportunidades de desarrollo del prototipo")
    para(
        doc,
        "AgroShield se diferencia porque no vende mensajería ni un tablero genérico; vende una señal Copernicus explicable para decisiones agrícolas bajo incertidumbre hídrica. Sus usuarios iniciales son técnicos de cuenca y equipos que asesoran productores. El valor se concentra en reducir ceguera operativa, priorizar muestreo, justificar pausas preventivas de riego y documentar evidencia para conversación institucional."
    )
    para(
        doc,
        "La escalabilidad debe ser progresiva: un AOI, luego varios tramos de la cuenca, luego otros nodos críticos. CDSE ofrece cuotas gratuitas útiles para el piloto [4], pero una expansión nacional exige control de frecuencia, cache y monitoreo de costo. AWS u otra nube se considera infraestructura futura, no requisito para demostrar el concepto."
    )

    heading(doc, "6.", "Conclusiones")
    para(
        doc,
        "AgroShield puede competir si se presenta como una respuesta simple y verificable: una cuenca real, una crisis real, un AOI real, una señal Sentinel real y una decisión agrícola simple. La contribución principal es convertir Copernicus en un semáforo trazable de inocuidad hídrica agrícola, sin ocultar incertidumbre ni prometer detección química directa."
    )
    para(
        doc,
        "El siguiente hito es ejecutar la prueba histórica y producir una primera evidencia visual. Si la señal se confirma, el MVP avanza hacia reporte reproducible y canal comunitario; si no, el proyecto conserva valor como atlas de riesgo hídrico o monitoreo ribereño. En ambos casos, la metodología evita sobreingeniería y mantiene foco en seguridad alimentaria."
    )

    unnumbered(doc, "Agradecimientos")
    para(
        doc,
        "Se reconoce el marco de innovación abierta del CopernicusLAC Hackathon 2026 y la guía metodológica de prototipado de la Jornada de Iniciación Científica de la Universidad Tecnológica de Panamá.",
        first_indent=False,
    )

    unnumbered(doc, "Referencias")
    refs = [
        "[1] CopernicusLAC Panama Centre, \"CopernicusLAC Hackathon 2026: Innovation in sustainable agriculture using Earth observation data,\" 2026. Available: https://www.copernicuslac-panama.eu/events-and-trainings/events-en/copernicus-lac-hackathon-2026-innovation-in-sustainable-agriculture-using-earth-observation-data/",
        "[2] European Space Agency, \"Sentinel-2 mission,\" 2026. Available: https://www.esa.int/Applications/Observing_the_Earth/Copernicus/Sentinel-2",
        "[3] European Space Agency, \"Sentinel-1 mission,\" 2026. Available: https://www.esa.int/Applications/Observing_the_Earth/Copernicus/Sentinel-1",
        "[4] Copernicus Data Space Ecosystem, \"Quotas and limitations,\" 2026. Available: https://documentation.dataspace.copernicus.eu/Quotas.html",
        "[5] Climate Hazards Center, University of California Santa Barbara, \"CHIRPS,\" 2026. Available: https://www.chc.ucsb.edu/data/chirps",
        "[6] H. Xu, \"Modification of normalised difference water index (NDWI) to enhance open water features in remotely sensed imagery,\" International Journal of Remote Sensing, vol. 27, no. 14, pp. 3025-3033, 2006.",
        "[7] S. Mishra and D. R. Mishra, \"Normalized difference chlorophyll index,\" Remote Sensing of Environment, vol. 117, pp. 394-406, 2012.",
        "[8] Ministerio de Ambiente de Panamá, \"MiAmbiente anuncia estabilización de la situación del río La Villa,\" 2025. Available: https://miambiente.gob.pa/miambiente-anuncia-estabilizacion-de-la-situacion-del-rio-la-villa-y-mejora-en-la-calidad-del-agua/",
        "[9] Ministerio de Desarrollo Agropecuario de Panamá, \"Avanzan acciones interinstitucionales para la recuperación de la cuenca del río La Villa,\" 2026. Available: https://mida.gob.pa/2026/01/19/avanzan-acciones-interinstitucionales-para-la-recuperacion-de-la-cuenca-del-rio-la-villa/",
        "[10] GLOBALG.A.P., \"Integrated Farm Assurance Version 6,\" 2026. Available: https://www.globalgap.org/uk_en/for-producers/globalg.a.p./integrated-farm-assurance-ifa/IFA-V6/",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        set_run_font(r, size=7.5)

    doc.save(JIC_OUT)


def h1(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=15, bold=True, color=ACCENT)


def h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True)


def build_roadmap():
    doc = Document()
    set_page(doc.sections[0])
    set_base_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("AgroShield")
    set_run_font(r, size=24, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Roadmap técnico, arquitectura y alcance modular")
    set_run_font(r, size=16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CopernicusLAC Hackathon 2026 - Documento interno de diseño")
    set_run_font(r, size=11, italic=True)

    h1(doc, "1. Decisión de producto")
    para(
        doc,
        "La versión competitiva de AgroShield debe iniciar como semáforo Copernicus de riesgo hídrico para inocuidad agrícola en Azuero. El producto no es una suite climática, ni un bot de mensajería, ni una plataforma nacional desde el primer día. El núcleo es una señal reproducible: AOI Chitré-La Arena, CDSE, Sentinel-2, línea base histórica, reglas transparentes y reporte técnico."
    )
    para(
        doc,
        "El usuario primario es el técnico de cuenca. Los productores, instituciones y comunidades reciben valor a través de reportes y canales secundarios. Esta separación evita que el MVP intente satisfacer simultáneamente fiscalización, exportación, comunicación pública y operación de planta."
    )

    h1(doc, "2. Alcance por fases")
    compact_table(
        doc,
        [
            ["Fase", "Objetivo", "Entregable", "Criterio de salida"],
            ["0. Spike interno", "Confirmar señal histórica.", "Script local + capturas + tabla.", "AOI y baseline defendibles."],
            ["1. MVP hackathon", "Semáforo reproducible.", "Pantalla/reporte con prueba reina.", "Estado explicado en 15 segundos."],
            ["2. Prototipo funcional", "Automatizar consulta y reporte.", "Pipeline CDSE + PDF/HTML.", "Resultados cacheados y auditables."],
            ["3. Canal comunitario", "Consulta y difusión.", "Bot WhatsApp secundario.", "No depende del bot para crear valor."],
            ["4. Escala cuenca", "Varios tramos Río La Villa.", "Config AOI multi-nodo.", "Costo y cuotas monitoreadas."],
            ["5. Módulos futuros", "Agronomía y expansión.", "NitroSync, ERA5-Land, más cuencas.", "Calibración con expertos."],
        ],
        [Cm(2.4), Cm(4.0), Cm(4.3), Cm(4.4)],
        header_fill=BLUE,
        font_size=8.2,
    )

    h1(doc, "3. Arquitectura del MVP")
    para(
        doc,
        "El MVP se ejecuta localmente o en una función serverless simple. La prioridad es reproducibilidad, no despliegue. Cada corrida debe tomar una configuración AOI, una fecha o ventana temporal, consultar o leer cache de CDSE, calcular señales y producir un objeto de evidencia."
    )
    compact_table(
        doc,
        [
            ["Capa", "Responsabilidad", "Implementación mínima", "Evolución"],
            ["Configuración", "AOI, fechas, umbrales.", "JSON/YAML local.", "Panel administrativo."],
            ["Adquisición", "CDSE Sentinel-2.", "Statistical API o cache.", "Scheduler + control de cuota."],
            ["Procesamiento", "Máscara agua, señal óptica, confianza.", "Python puro.", "Lambda/container."],
            ["Reglas", "Semáforo transparente.", "Condiciones if/else auditables.", "Calibración estadística."],
            ["Reporte", "Evidencia y recomendación.", "HTML/PDF simple.", "Plantillas institucionales."],
            ["Canal", "Entrega de reporte.", "Manual o enlace.", "WhatsApp comunitario."],
        ],
        [Cm(2.6), Cm(4.0), Cm(4.2), Cm(4.3)],
        header_fill=GREEN,
        font_size=8.2,
    )

    h2(doc, "Objeto de evidencia")
    para(
        doc,
        "El contrato interno debe ser pequeño. Un registro por observación contiene: fecha, AOI, fuente, porcentaje de nube, píxeles válidos, estadísticos de índice, baseline usado, estado, confianza, explicación y recomendaciones. Este objeto alimenta pantalla, reporte e interpretación asistida sin duplicar lógica."
    )
    compact_table(
        doc,
        [
            ["Campo", "Tipo", "Notas"],
            ["aoi_id", "string", "Ej. chitre_la_arena."],
            ["date", "YYYY-MM-DD", "Fecha de adquisición o ventana."],
            ["valid_pixels", "number", "Usado para confianza."],
            ["optical_signal", "number", "NDTI/red-green/MNDWI según validación."],
            ["baseline_delta", "number", "Diferencia contra control comparable."],
            ["status", "enum", "green, amber, red, uncertain."],
            ["explanation", "string", "Texto generado por plantilla/IA controlada."],
        ],
        [Cm(3.0), Cm(3.0), Cm(8.4)],
        header_fill=GRAY,
        font_size=8.2,
    )

    h1(doc, "4. Reglas del semáforo")
    para(
        doc,
        "No usar pesos compuestos hasta tener calibración. La primera versión debe operar con reglas que un jurado pueda auditar visualmente. El orden importa: primero confianza, luego agua, luego anomalía, luego contexto de lluvia."
    )
    compact_table(
        doc,
        [
            ["Orden", "Regla", "Resultado"],
            ["1", "Nube alta o píxeles válidos insuficientes.", "UNCERTAIN."],
            ["2", "No se confirma agua en el AOI.", "UNCERTAIN o AOI inválido."],
            ["3", "Señal óptica similar a baseline.", "GREEN."],
            ["4", "Señal óptica elevada vs. baseline.", "AMBER."],
            ["5", "Señal elevada + lluvia antecedente fuerte.", "RED."],
            ["6", "NDCI elevado.", "Flag biológico, no diagnóstico."],
        ],
        [Cm(1.7), Cm(8.2), Cm(5.0)],
        header_fill=YELLOW,
        font_size=8.2,
    )

    h1(doc, "5. Roadmap de módulos")
    h2(doc, "5.1 Reporte técnico automático")
    para(
        doc,
        "Debe ser la pieza de confianza. Contiene mapa o captura del AOI, fechas comparadas, señales, confianza, estado y limitación explícita. No se presenta como certificado GLOBALG.A.P.; se presenta como evidencia técnica para orientar decisión y priorizar muestreo."
    )
    h2(doc, "5.2 Interpretación asistida por IA")
    para(
        doc,
        "La IA no clasifica riesgo ni inventa datos. Recibe un objeto de evidencia y redacta un resumen controlado: qué se observó, qué tan confiable es, qué limita la lectura y qué acción preventiva se sugiere. En implementación, la salida debe pasar por plantilla con campos obligatorios para evitar alucinación."
    )
    h2(doc, "5.3 Bot comunitario WhatsApp")
    para(
        doc,
        "El bot es canal, no producto. Sirve para consultar estado, recibir enlace al reporte y eventualmente recolectar observaciones comunitarias. No debe depender de Twilio ni de despliegue real para el primer hito; puede simularse con mensajes generados desde el reporte."
    )
    h2(doc, "5.4 NitroSync")
    para(
        doc,
        "NitroSync queda como módulo futuro de seguridad alimentaria. Su idea es cruzar estado hídrico con condiciones agroclimáticas para orientar ventanas de fertilización. Requiere cultivo, etapa fenológica, práctica agronómica y validación con IDIAP/MIDA o asesores agrícolas. No pertenece al MVP corto."
    )
    h2(doc, "5.5 Escalamiento")
    para(
        doc,
        "La expansión debe hacerse por configuración, no por reescritura de código: lista de AOI, parámetros por tramo, frecuencia y cache. La arquitectura nacional solo es defendible después de probar un nodo."
    )

    h1(doc, "6. Arquitectura futura AWS")
    para(
        doc,
        "AWS se mantiene como arquitectura futura por afinidad del equipo, no como promesa de costo cero. Desde 2025, cuentas nuevas operan bajo un modelo de créditos y plan gratuito limitado; por tanto, el documento público debe evitar prometer doce meses universales gratis [12]."
    )
    compact_table(
        doc,
        [
            ["Servicio", "Uso", "Riesgo", "Mitigación"],
            ["Lambda", "Procesar AOI programado.", "Timeout o dependencias geoespaciales.", "Mantener payload liviano; CDSE procesa raster."],
            ["S3", "Guardar reportes.", "Costo bajo pero no cero eterno.", "Lifecycle y limpieza."],
            ["DynamoDB", "Historial de observaciones.", "Diseño de claves pobre.", "PK aoi_id, SK date."],
            ["EventBridge", "Programar corridas.", "Frecuencia excesiva.", "Cron alineado a disponibilidad Sentinel."],
            ["Secrets", "Credenciales CDSE.", "Costo/rotación.", "SSM/variables para demo, Secrets en producción."],
        ],
        [Cm(2.8), Cm(4.0), Cm(4.0), Cm(4.6)],
        header_fill=BLUE,
        font_size=8.1,
    )

    h1(doc, "7. Backlog de implementación")
    compact_table(
        doc,
        [
            ["Prioridad", "Historia", "Hecho cuando"],
            ["P0", "Como técnico, defino AOI y fechas.", "El sistema lee config y valida geometría."],
            ["P0", "Como técnico, comparo evento contra baseline.", "Se produce tabla con delta y confianza."],
            ["P0", "Como jurado, veo el semáforo.", "Estado explicado sin abrir código."],
            ["P1", "Como equipo, genero reporte.", "PDF/HTML reproduce la evidencia."],
            ["P1", "Como productor, recibo estado.", "Bot o mensaje muestra resumen."],
            ["P2", "Como agrónomo, evalúo NitroSync.", "Se documentan supuestos por cultivo."],
        ],
        [Cm(2.2), Cm(6.1), Cm(6.8)],
        header_fill=GREEN,
        font_size=8.1,
    )

    h1(doc, "8. Riesgos y decisiones")
    compact_table(
        doc,
        [
            ["Riesgo", "Impacto", "Decisión"],
            ["Río estrecho.", "Pocos píxeles válidos.", "Ampliar tramo o cambiar a atlas de riesgo."],
            ["Nubosidad.", "Sin señal óptica.", "Reportar incertidumbre y usar CHIRPS/S1 como contexto."],
            ["Química invisible.", "Claim incorrecto.", "Declarar proxy hidrosedimentario, no detección química."],
            ["Módulos excesivos.", "Pitch disperso.", "MVP solo Sentinel-2 + reglas + reporte."],
            ["IA decorativa.", "Preguntas de alucinación.", "IA como redacción controlada, no modelo de riesgo."],
            ["Costo nube.", "Promesa frágil.", "Local primero; nube como fase futura."],
        ],
        [Cm(4.0), Cm(4.3), Cm(6.8)],
        header_fill=RED,
        font_size=8.1,
    )

    h1(doc, "9. Entregables por fecha")
    para(
        doc,
        "Antes del inicio interno del hackathon se debe obtener la primera respuesta técnica: AOI, tres ventanas históricas, señal óptica, confianza y decisión de continuar o pivotar. Durante el hackathon, el esfuerzo público se concentra en contar la historia con evidencia y producir un MVP reproducible."
    )
    compact_table(
        doc,
        [
            ["Día", "Entrega"],
            ["D1", "AOI validado y fechas de prueba definidas."],
            ["D2", "Consulta CDSE o cache de Sentinel-2."],
            ["D3", "Baseline y comparación histórica."],
            ["D4", "Reglas de semáforo y primer reporte."],
            ["D5", "Pantalla/demo simple."],
            ["D6", "Narrativa, pitch y límites científicos."],
            ["D7", "Ensayo y respaldo técnico."],
            ["D8", "Paquete de propuesta listo."],
        ],
        [Cm(2.0), Cm(13.0)],
        header_fill=GRAY,
        font_size=8.2,
    )

    h1(doc, "10. Diseño de datos")
    para(
        doc,
        "El diseño de datos debe ser deliberadamente pequeño para que el equipo pueda depurar el sistema durante la hackathon. La fuente de verdad no es el PDF ni la pantalla, sino un conjunto de registros JSON que puedan reconstruir cada decisión. Esto permite repetir la demo, explicar resultados al jurado y migrar luego a DynamoDB sin cambiar el modelo mental."
    )
    compact_table(
        doc,
        [
            ["Archivo/colección", "Contenido", "Uso"],
            ["aoi_config.json", "Nombre, geometría, notas hidrológicas, versión.", "Define el tramo evaluado."],
            ["observations.jsonl", "Una observación por fecha y fuente.", "Auditoría y comparación histórica."],
            ["baselines.json", "Ventanas comparables y estadísticos.", "Evitar umbrales inventados."],
            ["alerts.jsonl", "Estado, confianza, reglas activadas.", "Entrada de reporte y demo."],
            ["report_manifest.json", "Fuentes, imágenes, tablas y hashes.", "Reproducir el PDF/HTML."],
        ],
        [Cm(4.0), Cm(6.0), Cm(5.0)],
        header_fill=BLUE,
        font_size=8.1,
    )
    h2(doc, "10.1 Esquema de observación")
    compact_table(
        doc,
        [
            ["Campo", "Ejemplo", "Regla"],
            ["source", "sentinel2_l2a", "No mezclar fuentes en un mismo registro."],
            ["window", "2025-06-10/2025-06-12", "Siempre guardar rango, no solo fecha."],
            ["cloud_pct", "18.4", "Si supera umbral, estado puede ser incierto."],
            ["water_pixels", "42", "Mide si el AOI es útil."],
            ["index_mean", "0.137", "Guardar media y percentiles."],
            ["baseline_ref", "baseline_2025_wet", "Toda anomalía debe citar baseline."],
            ["quality_flags", "cloud_edge,mixed_pixels", "Explicar incertidumbre."],
        ],
        [Cm(3.5), Cm(4.0), Cm(7.4)],
        header_fill=GRAY,
        font_size=8.1,
    )

    h1(doc, "11. Estrategia de AOI y fallback")
    para(
        doc,
        "El mayor riesgo técnico es la geometría del río. Un AOI demasiado pequeño puede producir pocos píxeles válidos; uno demasiado grande puede mezclar agua, suelo, urbano y vegetación. La estrategia correcta es probar tres geometrías y escoger la que preserve señal de agua sin perder contexto."
    )
    compact_table(
        doc,
        [
            ["AOI", "Ventaja", "Riesgo", "Uso"],
            ["Nodo estrecho", "Más cercano a la toma.", "Pocos píxeles; mezcla espectral.", "Solo si datos son limpios."],
            ["Tramo ampliado", "Más píxeles y mejor estadística.", "Menos precisión local.", "Opción recomendada para MVP."],
            ["Buffer de cuenca baja", "Más contexto hidrológico.", "Puede diluir la señal.", "Fallback para atlas de riesgo."],
        ],
        [Cm(3.3), Cm(4.0), Cm(4.0), Cm(3.8)],
        header_fill=YELLOW,
        font_size=8.1,
    )
    para(
        doc,
        "La decisión de AOI no debe tomarse por estética del mapa. Debe tomarse por número de píxeles válidos, estabilidad de baseline y capacidad de explicar la señal. Si el nodo exacto falla, no se abandona el proyecto: se cambia a tramo ampliado o a monitoreo de riesgo de cuenca."
    )

    h1(doc, "12. Pipeline técnico detallado")
    compact_table(
        doc,
        [
            ["Paso", "Entrada", "Proceso", "Salida"],
            ["1", "AOI + fechas", "Validar geometría y ventana.", "request_plan.json"],
            ["2", "CDSE", "Consultar Sentinel-2 L2A.", "raw_stats.json"],
            ["3", "Stats", "Filtrar nubes y píxeles inválidos.", "observation.json"],
            ["4", "Baseline", "Comparar con ventana equivalente.", "delta.json"],
            ["5", "Reglas", "Aplicar semáforo transparente.", "alert.json"],
            ["6", "Reporte", "Renderizar tablas y explicación.", "report.html/pdf"],
            ["7", "Demo", "Mostrar mapa, fechas y estado.", "pantalla/pitch"],
        ],
        [Cm(1.2), Cm(3.0), Cm(6.4), Cm(4.3)],
        header_fill=GREEN,
        font_size=8.0,
    )
    h2(doc, "12.1 Estructura sugerida del repositorio")
    compact_table(
        doc,
        [
            ["Ruta", "Responsabilidad"],
            ["configs/aoi/", "GeoJSON y metadatos de Chitré-La Arena."],
            ["data/cache/", "Respuestas CDSE y datasets descargados."],
            ["src/ingest/", "Clientes CDSE/CHIRPS y normalización."],
            ["src/rules/", "Reglas de semáforo y confianza."],
            ["src/report/", "Plantillas HTML/PDF y manifest."],
            ["app/", "Demo local o Streamlit ligero."],
            ["docs/evidence/", "Capturas y tablas para pitch."],
        ],
        [Cm(5.0), Cm(10.0)],
        header_fill=GRAY,
        font_size=8.1,
    )

    h1(doc, "13. Diseño del reporte técnico")
    para(
        doc,
        "El reporte es el artefacto de confianza. Debe poder leerse sin conocer el código. No conviene llamarlo certificado, dictamen legal ni auditoría GLOBALG.A.P.; debe llamarse reporte técnico de evidencia satelital para decisión preventiva."
    )
    compact_table(
        doc,
        [
            ["Sección", "Contenido obligatorio"],
            ["Resumen", "Estado, fecha, AOI y recomendación."],
            ["Evidencia", "Tabla de evento vs. baseline y control."],
            ["Confianza", "Nube, píxeles válidos, fuente, limitaciones."],
            ["Interpretación", "Texto controlado a partir del alert.json."],
            ["Acción sugerida", "Revisar riego, muestreo, inspección o espera."],
            ["Trazabilidad", "Fuentes, parámetros y versión del modelo."],
        ],
        [Cm(4.0), Cm(10.9)],
        header_fill=BLUE,
        font_size=8.1,
    )
    h2(doc, "13.1 Guardrails para IA")
    compact_table(
        doc,
        [
            ["Regla", "Motivo"],
            ["La IA no calcula el estado.", "Evita decisiones opacas."],
            ["La IA solo reescribe campos del objeto de evidencia.", "Evita alucinación."],
            ["Toda frase de incertidumbre es obligatoria.", "Evita sobreconfianza."],
            ["No mencionar pesticidas detectados.", "La señal no mide químicos disueltos."],
            ["No recomendar dosis agronómicas.", "Fuera del MVP y requiere calibración."],
        ],
        [Cm(5.5), Cm(9.5)],
        header_fill=YELLOW,
        font_size=8.1,
    )

    h1(doc, "14. Validación y pruebas")
    para(
        doc,
        "Las pruebas deben cubrir ciencia, software y narrativa. No basta con que el código corra: debe poder explicar por qué el estado es verde, amarillo, rojo o incierto. Cada prueba debe guardar evidencia para el pitch."
    )
    compact_table(
        doc,
        [
            ["Tipo", "Prueba", "Aceptación"],
            ["Datos", "CDSE devuelve estadísticas para el AOI.", "Sin errores y con píxeles válidos."],
            ["Geometría", "Comparar tres AOI.", "Se escoge el más estable."],
            ["Baseline", "Evento vs. control.", "Diferencia visible o incertidumbre honesta."],
            ["Reglas", "Casos synthetic de nube/píxeles.", "Nunca fuerza verde con baja confianza."],
            ["Reporte", "Reconstrucción desde JSON.", "No depende de edición manual."],
            ["Pitch", "Explicación en 60 segundos.", "Jurado entiende decisión y límite."],
        ],
        [Cm(2.6), Cm(6.0), Cm(6.4)],
        header_fill=GREEN,
        font_size=8.1,
    )

    h1(doc, "15. Operación, costos y seguridad")
    para(
        doc,
        "Para el hackathon, la operación recomendada es local-first. Se evita dependencia de cuentas, dominios, permisos, SMS y costos. En una fase posterior, la nube debe entrar con presupuestos, alarmas y límites de frecuencia. Las credenciales CDSE no deben quedar en el repositorio ni en el documento público."
    )
    compact_table(
        doc,
        [
            ["Tema", "Decisión"],
            ["Costos", "Usar cache y pocas fechas; documentar cuotas CDSE."],
            ["Credenciales", "Variables de entorno o archivo local no versionado."],
            ["Privacidad", "No recolectar datos personales en el MVP."],
            ["WhatsApp", "Canal comunitario opt-in, fase posterior."],
            ["Licencias", "Citar fuentes y mantener trazabilidad de datos."],
            ["Resiliencia", "Demo debe funcionar con datos cacheados si CDSE falla."],
        ],
        [Cm(4.0), Cm(11.0)],
        header_fill=GRAY,
        font_size=8.1,
    )

    h1(doc, "16. Matriz de decisión de pivot")
    para(
        doc,
        "El proyecto debe entrar al hackathon con pivots preaprobados. Esto evita perder tiempo defendiendo una señal débil. Todos los pivots conservan el activo principal: evidencia Copernicus para seguridad alimentaria en Azuero."
    )
    compact_table(
        doc,
        [
            ["Resultado del spike", "Decisión"],
            ["Señal óptica clara en AOI.", "Continuar con semáforo hídrico agrícola."],
            ["Señal existe solo en tramo amplio.", "Cambiar narrativa a tramo de cuenca baja."],
            ["Nube impide prueba óptica.", "Priorizar CHIRPS/Sentinel-1 como alerta contextual."],
            ["Río demasiado estrecho.", "Pivot a atlas de riesgo de cuenca o monitoreo ribereño."],
            ["No hay anomalía histórica.", "Presentar herramienta de verificación y límites."],
        ],
        [Cm(6.5), Cm(8.5)],
        header_fill=RED,
        font_size=8.1,
    )

    h1(doc, "17. Referencias base")
    refs = [
        "[1] CopernicusLAC Panama Centre, \"CopernicusLAC Hackathon 2026,\" 2026.",
        "[2] European Space Agency, \"Sentinel-2 mission,\" 2026.",
        "[3] European Space Agency, \"Sentinel-1 mission,\" 2026.",
        "[4] Copernicus Data Space Ecosystem, \"Quotas and limitations,\" 2026.",
        "[5] Climate Hazards Center, \"CHIRPS,\" University of California Santa Barbara, 2026.",
        "[6] ECMWF, \"ERA5-Land hourly data,\" Copernicus Climate Data Store, 2026.",
        "[7] H. Xu, \"Modification of normalised difference water index,\" International Journal of Remote Sensing, 2006.",
        "[8] S. Mishra and D. R. Mishra, \"Normalized difference chlorophyll index,\" Remote Sensing of Environment, 2012.",
        "[9] Ministerio de Ambiente de Panamá, \"Situación del río La Villa,\" 2025.",
        "[10] Ministerio de Desarrollo Agropecuario de Panamá, \"Recuperación de la cuenca del río La Villa,\" 2026.",
        "[11] GLOBALG.A.P., \"Integrated Farm Assurance Version 6,\" 2026.",
        "[12] Amazon Web Services, \"AWS Free Tier,\" 2026.",
    ]
    for ref in refs:
        para(doc, ref, size=9, first_indent=False, after=1)

    doc.save(ROADMAP_OUT)


if __name__ == "__main__":
    build_jic_v2()
    build_roadmap()
    print(JIC_OUT)
    print(ROADMAP_OUT)
