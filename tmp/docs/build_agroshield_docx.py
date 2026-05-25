from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\jose\Documents\CopernicusLAC")
OUT = ROOT / "AgroShield_Propuesta_Tecnica_JIC_UTP.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=8.0, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_columns(section, num=2, space_twips=567):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if not cols.getparent():
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def set_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def set_run_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", lead=None, size=10, first_indent=True, justify=True, after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if lead:
        r = p.add_run(lead + " ")
        set_run_font(r, size=size, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def add_heading(doc, number, title, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f"{number} {title}")
    set_run_font(r, size=14 if level == 1 else 10, bold=True)
    return p


def add_unnumbered_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    set_run_font(r, size=14, bold=True)
    return p


def add_caption(doc, text, kind="table"):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "table" else WD_ALIGN_PARAGRAPH.JUSTIFY
    first, rest = text.split(" ", 1)
    r1 = p.add_run(first + " ")
    set_run_font(r1, size=8, bold=True)
    r2 = p.add_run(rest)
    set_run_font(r2, size=8)
    return p


def add_source(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run("Fuente: ")
    set_run_font(r1, size=8, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_compact_table(doc, rows, widths, header_fill="D9EAD3"):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, val in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], val, bold=True, size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], header_fill)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=7.6)
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    set_page(sec)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    for section in doc.sections:
        section.different_first_page_header_footer = False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Categoría prototipado: seguridad alimentaria y agricultura sostenible con observación de la Tierra.")
    set_run_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("AgroShield: sistema satelital de evidencia temprana para riesgo hídrico agrícola en la cuenca del río La Villa")
    set_run_font(r, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("AgroShield: satellite evidence system for early agricultural water-risk assessment in the La Villa river basin")
    set_run_font(r, size=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("José Rodríguez1*, Yanesis Valdés2")
    set_run_font(r, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("1Afiliación pendiente, grupo o unidad académica pendiente; 2Afiliación pendiente, grupo o unidad académica pendiente")
    set_run_font(r, size=10, italic=True)

    add_para(
        doc,
        "AgroShield propone un prototipo de apoyo a decisiones para técnicos de cuenca que necesitan anticipar riesgo hídrico agrícola en la cuenca del río La Villa, Azuero, Panamá. La solución integra datos Copernicus Sentinel-2 y Sentinel-1 con lluvia CHIRPS y variables agroclimáticas ERA5-Land para generar un semáforo técnico, un reporte PDF automático y una interpretación asistida por IA. El sistema no pretende detectar pesticidas disueltos ni reemplazar análisis de laboratorio; su aporte es organizar evidencia satelital e hidrometeorológica asociada con sedimento, clorofila, saturación antecedente y condiciones de manejo agrícola. El prototipo se encuentra en fase inicial y será verificado internamente mediante fechas históricas de 2025 antes de construir el entregable formal del hackathon.",
        lead="Resumen",
        first_indent=False,
    )
    add_para(
        doc,
        "AgroShield, Copernicus, riesgo hídrico, seguridad alimentaria, Sentinel.",
        lead="Palabras clave",
        first_indent=False,
    )
    add_para(
        doc,
        "AgroShield is an early-stage decision-support prototype for watershed technicians assessing agricultural water risk in the La Villa river basin, Panama. It combines Copernicus Sentinel-2 and Sentinel-1 observations with CHIRPS rainfall and ERA5-Land agroclimatic variables to produce a technical traffic-light status, an automated PDF evidence report, and AI-assisted interpretation. The prototype does not claim direct detection of dissolved pesticides; it structures satellite and hydrometeorological evidence related to sediment transport, chlorophyll, antecedent saturation and agricultural decision windows.",
        lead="Abstract",
        first_indent=False,
    )
    add_para(
        doc,
        "AgroShield, Copernicus, food security, Sentinel, water risk.",
        lead="Keywords",
        first_indent=False,
    )
    p = doc.add_paragraph()
    r = p.add_run("* Corresponding author: jose.rodriguez104@utp.ac.pa")
    set_run_font(r, size=8)

    body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(body_sec)
    set_columns(body_sec, 2)

    add_heading(doc, "1.", "Introducción")
    add_para(
        doc,
        "La seguridad alimentaria depende de la calidad y disponibilidad del agua usada para consumo humano, riego, lavado y continuidad de cadenas agroexportadoras. En Azuero, la cuenca del río La Villa concentra valor productivo y vulnerabilidad ambiental: eventos de lluvia intensa pueden movilizar sedimento y contaminantes desde fuentes difusas hacia tomas y zonas agrícolas. La observación de la Tierra permite observar la cuenca completa con una continuidad imposible de lograr solo con visitas de campo o sensores puntuales [1]-[4]."
    )
    add_para(
        doc,
        "El problema de prototipado consiste en convertir datos satelitales abiertos en evidencia accionable para un técnico de cuenca. Las restricciones principales son costo bajo, trazabilidad, lenguaje comprensible, límites científicos explícitos y capacidad de validar el flujo con datos históricos antes de operar. A diferencia de un tablero genérico, AgroShield prioriza un reporte técnico que explique qué se observó, qué tan confiable es la observación y qué decisión preventiva podría considerarse."
    )

    add_heading(doc, "2.", "Diseño y metodología")
    add_para(
        doc,
        "El diseño se organiza como un motor de evidencia. Primero se define un área de interés en el corredor Chitré-La Arena del río La Villa. Luego se consultan productos Sentinel-2 Level-2A para señales ópticas, Sentinel-1 para continuidad bajo nubosidad, CHIRPS para lluvia antecedente y ERA5-Land para condiciones agroclimáticas complementarias [2]-[6]. La salida se resume en componentes separados antes de calcular un índice compuesto preliminar."
    )

    add_caption(doc, "Tabla 1. Ficha técnica del prototipo")
    add_compact_table(
        doc,
        [
            ["Elemento", "Definición operativa"],
            ["Usuario", "Técnico de cuenca que interpreta evidencia y coordina validación local."],
            ["Zona piloto", "Río La Villa, Azuero, Panamá; nodo Chitré-La Arena."],
            ["Datos base", "Sentinel-2, Sentinel-1, CDSE, CHIRPS y ERA5-Land."],
            ["Salida", "Semáforo técnico, PDF automático, resumen IA y canal comunitario."],
            ["Estado", "Fase inicial; código interno descartable para validar factibilidad."],
        ],
        [Cm(2.0), Cm(6.0)],
    )
    add_source(doc, "Elaboración propia con base en el alcance del prototipo.")

    add_para(
        doc,
        "El puntaje no se presenta como una caja negra. Se separan cuatro señales: riesgo óptico, riesgo biológico, riesgo hidrometeorológico y confianza de observación. La confianza evita clasificar como verde cuando existen nubes, pocos píxeles válidos o ausencia de datos. El índice compuesto preliminar se expresa como CRI = 0.45RO + 0.20RB + 0.25RH + 0.10C, donde RO es riesgo óptico, RB riesgo biológico, RH riesgo hidrometeorológico y C confianza invertida o penalización de incertidumbre. Los pesos son iniciales y deben calibrarse con datos locales."
    )

    add_caption(doc, "Tabla 2. Componentes auditables del semáforo técnico")
    add_compact_table(
        doc,
        [
            ["Componente", "Evidencia usada"],
            ["RO", "MNDWI/NDTI o reflectancia visible para agua y sedimento."],
            ["RB", "NDCI como señal exploratoria de clorofila y eutrofización."],
            ["RH", "CHIRPS y Sentinel-1 para lluvia, saturación y continuidad."],
            ["C", "Nubosidad, píxeles válidos, fecha y resolución espacial."],
            ["CRI", "Score preliminar; no sustituye análisis de laboratorio."],
        ],
        [Cm(2.0), Cm(6.0)],
    )
    add_source(doc, "Elaboración propia a partir de índices de literatura y fuentes abiertas.")

    add_heading(doc, "2.1", "Módulos del prototipo", level=2)
    add_para(
        doc,
        "El módulo Evidence Engine consulta o cachea observaciones satelitales y produce registros simples. El generador PDF compila fecha, área, fuentes, valores y limitaciones como evidencia técnica, no como certificación oficial. La capa de IA redacta una interpretación breve basada únicamente en datos calculados: estado, confianza, razón y recomendación. El bot comunitario de WhatsApp queda como canal secundario para consultar estado o recibir reportes, sin convertirse en el producto principal."
    )
    add_para(
        doc,
        "NitroSync se mantiene como módulo de seguridad alimentaria porque vincula el riesgo hídrico con decisiones de fertilización. Cuando el estado hídrico es favorable y ERA5-Land sugiere condiciones de radiación adecuadas, el sistema puede recomendar una ventana de manejo agronómico. En esta fase se documenta como módulo complementario sujeto a calibración con especialistas agrícolas."
    )

    add_heading(doc, "3.", "Resultados esperados")
    add_para(
        doc,
        "El resultado esperado de la primera etapa es una propuesta técnicamente verificable y un prototipo interno descartable que demuestre si existe señal suficiente en fechas históricas. Se espera obtener una tabla comparativa entre días de crisis, línea base cercana y control estacional, junto con capturas o valores que permitan defender el uso de Copernicus en el caso del río La Villa."
    )
    add_caption(doc, "Tabla 3. Resultado mínimo de la verificación interna")
    add_compact_table(
        doc,
        [
            ["Prueba", "Criterio de aceptación"],
            ["Área", "El AOI contiene píxeles de agua suficientes o se redefine."],
            ["Óptico", "La señal Sentinel-2 no queda anulada por nubosidad."],
            ["Histórico", "La fecha de 2025 difiere de una línea base comparable."],
            ["Reporte", "El PDF explica estado, confianza, límites y fuentes."],
            ["IA", "El resumen no inventa datos y cita incertidumbre."],
        ],
        [Cm(2.0), Cm(6.0)],
    )
    add_source(doc, "Criterios internos propuestos para la semana previa al hackathon.")

    add_heading(doc, "4.", "Construcción y validación del prototipo")
    add_para(
        doc,
        "La construcción se realizará por etapas. En la etapa cero se usa código local temporal para consultar CDSE o datos cacheados y validar el área de interés. En la etapa uno se genera el semáforo técnico y el reporte PDF. En la etapa dos se agrega la interpretación asistida por IA y el canal comunitario. El código de prueba previo al hackathon puede eliminarse para evitar dependencia de implementaciones apresuradas; lo que debe permanecer es el aprendizaje técnico, las capturas y las decisiones de diseño."
    )
    add_para(
        doc,
        "La validación no debe prometer operación institucional. Para evitar sobreingeniería, el prototipo inicia con un nodo y un usuario principal. Si el río resulta demasiado estrecho para una señal óptica robusta, el plan alterno es ampliar el tramo de análisis, usar un atlas de riesgo de cuenca o priorizar monitoreo ribereño con Sentinel-1/Sentinel-2."
    )

    add_heading(doc, "5.", "Oportunidades de desarrollo del prototipo")
    add_para(
        doc,
        "El mercado inicial no es un consumidor masivo, sino equipos técnicos que requieren trazabilidad para decidir, comunicar y priorizar inspecciones. El valor frente a soluciones convencionales es que Copernicus aporta visión sinóptica, archivo histórico y bajo costo marginal [1], [4]. Las oportunidades incluyen asistencia a productores exportadores, soporte a programas de recuperación de cuencas, evidencia para inspección ambiental y módulos agronómicos como NitroSync."
    )
    add_para(
        doc,
        "Los costos reales dependen de cuotas, despliegue y frecuencia. Para 2026 debe evitarse prometer doce meses gratuitos universales en AWS, pues las cuentas nuevas operan con un modelo de créditos y plan gratuito limitado [12]. Una arquitectura futura puede usar funciones serverless, almacenamiento de reportes y base de datos liviana; sin embargo, el prototipo puede validarse localmente sin costo de nube."
    )

    add_heading(doc, "6.", "Conclusiones")
    add_para(
        doc,
        "AgroShield aporta una propuesta de investigación aplicada donde la innovación no es el envío de alertas, sino la conversión de evidencia satelital en una explicación técnica usable. Sus contribuciones son: separar ciencia de comunicación, reconocer límites de detección química, usar Copernicus para observar cuencas completas y producir reportes trazables para decisiones preventivas."
    )
    add_para(
        doc,
        "La principal limitación es la necesidad de validar resolución espacial, nubosidad y umbrales con datos reales de la cuenca. Aun así, el enfoque es pertinente para seguridad alimentaria porque conecta agua, inocuidad, manejo agronómico y resiliencia de productores. El trabajo futuro debe calibrar pesos del índice, incorporar muestreos locales y evaluar escalamiento a otros nodos críticos de Panamá."
    )

    add_unnumbered_heading(doc, "Agradecimientos")
    add_para(
        doc,
        "Se reconoce el marco de innovación abierta del CopernicusLAC Hackathon 2026 y la orientación metodológica del instructivo de prototipado de la Jornada de Iniciación Científica de la Universidad Tecnológica de Panamá.",
        first_indent=False,
    )

    add_unnumbered_heading(doc, "Referencias")
    references = [
        "[1] CopernicusLAC Panama Centre, \"CopernicusLAC Hackathon 2026: Innovation in sustainable agriculture using Earth observation data,\" 2026. Available: https://www.copernicuslac-panama.eu/events-and-trainings/events-en/copernicus-lac-hackathon-2026-innovation-in-sustainable-agriculture-using-earth-observation-data/",
        "[2] European Space Agency, \"Sentinel-2 mission,\" 2026. Available: https://www.esa.int/Applications/Observing_the_Earth/Copernicus/Sentinel-2",
        "[3] European Space Agency, \"Sentinel-1 mission,\" 2026. Available: https://www.esa.int/Applications/Observing_the_Earth/Copernicus/Sentinel-1",
        "[4] Copernicus Data Space Ecosystem, \"Quotas and limitations,\" 2026. Available: https://documentation.dataspace.copernicus.eu/Quotas.html",
        "[5] Climate Hazards Center, University of California Santa Barbara, \"CHIRPS: Climate Hazards Group InfraRed Precipitation with Station data,\" 2026. Available: https://www.chc.ucsb.edu/data/chirps",
        "[6] ECMWF, \"ERA5-Land hourly data from 1950 to present,\" Copernicus Climate Data Store, 2026. Available: https://cds.climate.copernicus.eu/",
        "[7] H. Xu, \"Modification of normalised difference water index (NDWI) to enhance open water features in remotely sensed imagery,\" International Journal of Remote Sensing, vol. 27, no. 14, pp. 3025-3033, 2006.",
        "[8] S. Mishra and D. R. Mishra, \"Normalized difference chlorophyll index: A novel model for remote estimation of chlorophyll-a concentration in turbid productive waters,\" Remote Sensing of Environment, vol. 117, pp. 394-406, 2012.",
        "[9] Ministerio de Ambiente de Panamá, \"MiAmbiente anuncia estabilización de la situación del río La Villa y mejora en la calidad del agua,\" 2025. Available: https://miambiente.gob.pa/miambiente-anuncia-estabilizacion-de-la-situacion-del-rio-la-villa-y-mejora-en-la-calidad-del-agua/",
        "[10] Ministerio de Desarrollo Agropecuario de Panamá, \"Avanzan acciones interinstitucionales para la recuperación de la cuenca del río La Villa,\" 2026. Available: https://mida.gob.pa/2026/01/19/avanzan-acciones-interinstitucionales-para-la-recuperacion-de-la-cuenca-del-rio-la-villa/",
        "[11] GLOBALG.A.P., \"Integrated Farm Assurance Version 6 for fruit and vegetables,\" 2026. Available: https://www.globalgap.org/uk_en/for-producers/globalg.a.p./integrated-farm-assurance-ifa/IFA-V6/",
        "[12] Amazon Web Services, \"AWS Free Tier,\" 2026. Available: https://docs.aws.amazon.com/awsaccountbilling/latest/aboutv2/free-tier.html",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        set_run_font(r, size=7.7)

    # Ensure all runs use Times New Roman, including table text.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
